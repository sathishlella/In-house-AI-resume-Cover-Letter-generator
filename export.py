from typing import Dict, List
from jinja2 import Template
from docx import Document
from docx.shared import Pt

def render_resume_md(data: Dict, jinja_template: str) -> str:
    tmpl = Template(jinja_template)
    return tmpl.render(**data)

def render_cover_md(data: Dict, jinja_template: str) -> str:
    tmpl = Template(jinja_template)
    return tmpl.render(**data)

def md_to_docx(md_text: str, title: str = "Document") -> Document:
    # Simple Markdown -> DOCX (headings, bullets). For richer: use mammoth/weasyprint pipeline.
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style=None).style = None
        else:
            doc.add_paragraph(line)
    return doc

def save_docx(doc: Document, filepath: str):
    doc.save(filepath)
